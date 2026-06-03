"""Document parsers for the Phase 2 Path A (prose-to-Rego) ingestion.

Five supported MIME types, sniffed from file bytes — the client's
declared Content-Type and the filename extension are advisory only:

    application/pdf                                    .pdf
    application/vnd.openxmlformats-officedocument.wordprocessingml.document
                                                       .docx
    text/markdown                                      .md
    text/html                                          .html
    text/plain                                         .txt

Each parser is called inside `asyncio.wait_for(timeout=…)` from the
router — malicious documents (zip bombs, recursive PDFs) cannot
hold up the event loop. Parsing runs in a worker thread because the
underlying libraries are sync + CPU-bound.

If a parser returns empty / whitespace-only text, the router refuses
the upload with 422. Scanned PDFs are the common case — there's
nothing the LLM step in PR 6 can do with an empty document.
"""
from __future__ import annotations

import asyncio
import io
from typing import Callable, NamedTuple

import pypdf
from bs4 import BeautifulSoup
from docx import Document


# Public MIME labels — what we store in `policy_uploads.sniffed_mime`.
MIME_PDF = "application/pdf"
MIME_DOCX = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
MIME_MARKDOWN = "text/markdown"
MIME_HTML = "text/html"
MIME_TEXT = "text/plain"


class UnsupportedFormat(ValueError):
    """Raised by `sniff_mime` when none of the supported formats match.
    Routers translate to 415."""


class EmptyExtraction(ValueError):
    """Raised by `extract_text` when the parser succeeded but produced
    nothing useful — typical for scanned/image-only PDFs. Routers
    translate to 422 with an OCR-required hint."""


class ParseTimeout(Exception):
    """Raised when a parser exceeds the configured timeout. Routers
    translate to 408."""


class _Sniff(NamedTuple):
    mime: str
    parser: Callable[[bytes], str]


def sniff_mime(raw: bytes, *, filename_hint: str | None = None) -> str:
    """Detect the file type from leading bytes.

    Filename hint disambiguates `text/markdown` vs `text/plain` (both
    are decoded the same way; the label affects how the LLM in PR 6
    is prompted). It is *not* used to bypass binary detection — a
    `.txt` extension with a PDF magic header is parsed as PDF.
    """
    if raw.startswith(b"%PDF-"):
        return MIME_PDF
    # DOCX is a zip; OOXML's first byte sequence is the PK zip header
    # followed (in the central directory) by an entry named "word/".
    # We check both to avoid mistaking a generic .zip for a .docx.
    if raw[:4] == b"PK\x03\x04" and b"word/" in raw[:4096]:
        return MIME_DOCX
    if raw[:4] == b"PK\x03\x04":
        raise UnsupportedFormat(
            "ZIP archive that is not a .docx. If this is a legacy .doc, "
            "open in Word and Save As .docx."
        )
    # Text path — decode-then-classify. UTF-8 first, fall back to
    # latin-1 (lossless single-byte) so a non-UTF-8 .md still parses.
    try:
        text_head = raw[:4096].decode("utf-8")
    except UnicodeDecodeError:
        try:
            text_head = raw[:4096].decode("latin-1")
        except UnicodeDecodeError:
            raise UnsupportedFormat("binary content of unknown type")
    stripped = text_head.lstrip()
    if stripped.startswith("<") and ("<html" in stripped.lower() or "<!doctype html" in stripped.lower()):
        return MIME_HTML
    # Markdown is hard to detect by bytes alone — fall back to the
    # filename hint here only. If the hint disagrees, plain text wins.
    if filename_hint and filename_hint.lower().endswith((".md", ".markdown")):
        return MIME_MARKDOWN
    return MIME_TEXT


def _parse_pdf(raw: bytes) -> str:
    reader = pypdf.PdfReader(io.BytesIO(raw))
    chunks: list[str] = []
    for page in reader.pages:
        try:
            chunks.append(page.extract_text() or "")
        except Exception:
            # Don't fail the whole document because page N has a quirk;
            # skip the page, surface what we can.
            chunks.append("")
    return "\n\n".join(c.strip() for c in chunks if c.strip())


def _parse_docx(raw: bytes) -> str:
    doc = Document(io.BytesIO(raw))
    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())
    # Tables — flatten cell-by-cell. Sequence within a table row is
    # preserved; rows are joined with newlines.
    for table in doc.tables:
        for row in table.rows:
            row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_cells:
                parts.append(" | ".join(row_cells))
    return "\n".join(parts)


def _parse_markdown(raw: bytes) -> str:
    # Markdown is its own plaintext — no rendering. The LLM in PR 6
    # is perfectly capable of reading raw markdown.
    return raw.decode("utf-8", errors="replace")


def _parse_html(raw: bytes) -> str:
    soup = BeautifulSoup(raw, "html.parser")
    # Remove script/style noise outright.
    for tag in soup(("script", "style", "noscript", "head")):
        tag.decompose()
    # `separator=\n` keeps paragraph structure that the LLM cares about.
    return soup.get_text(separator="\n", strip=True)


def _parse_text(raw: bytes) -> str:
    return raw.decode("utf-8", errors="replace")


_PARSERS: dict[str, Callable[[bytes], str]] = {
    MIME_PDF: _parse_pdf,
    MIME_DOCX: _parse_docx,
    MIME_MARKDOWN: _parse_markdown,
    MIME_HTML: _parse_html,
    MIME_TEXT: _parse_text,
}


async def extract_text(*, raw: bytes, mime: str, timeout_seconds: int) -> str:
    """Run the parser for `mime` with the given timeout.

    Raises:
        UnsupportedFormat — mime not in the supported set
        ParseTimeout      — parser exceeded the timeout
        EmptyExtraction   — parser produced empty/whitespace output
    """
    parser = _PARSERS.get(mime)
    if parser is None:
        raise UnsupportedFormat(f"no parser for {mime}")

    try:
        text = await asyncio.wait_for(
            asyncio.to_thread(parser, raw),
            timeout=timeout_seconds,
        )
    except asyncio.TimeoutError as exc:
        raise ParseTimeout(f"parser exceeded {timeout_seconds}s") from exc

    if not text or not text.strip():
        raise EmptyExtraction(
            "no extractable text — scanned image / OCR required"
        )
    return text
